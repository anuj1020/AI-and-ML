# import streamlit as st
import os

# To save the files to storage
def save_files(uploaded_files):
    path = os.getcwd()
    save_dir = path + "/uploaded_files"
    os.makedirs(save_dir, exist_ok=True)  # Create the directory if it doesn't exist
    
    for uploaded_file in uploaded_files:
        file_path = os.path.join(save_dir, uploaded_file.name)

        # Save the file
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

# Function to extract and return content from uploaded files
def getContentFromUploadedFiles(file_name):
    path = os.getcwd()
    file = path + "/uploaded_files/"+file_name
    
    if file_name.endswith('.pdf'):
        import PyPDF2
        with open(file, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            text = ''
            for page in reader.pages:
                text += page.extract_text() + '\n'
            contents = text

    elif file_name.endswith('.docx'):
        from docx import Document
        
        doc = Document(file_name)

        # Extract text
        for para in doc.paragraphs:
            print(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                print(row_data)

    elif file_name.endswith('.doc'):
        import textract
        text = textract.process(file).decode('utf-8')
        contents = text
    
    for content in contents:
        print(content)

    return contents
    
    